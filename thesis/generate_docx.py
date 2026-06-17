# -*- coding: utf-8 -*-
"""
generate_docx.py
================
Parses Thesis_TDT.md and generates a formatted Word document (.docx) 
complying with Ton Duc Thang University formatting guidelines.

Requirements:
    pip install python-docx
"""
import sys
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

# --- Constants & Paths ---
ROOT = Path(__file__).parents[1]
MD_PATH = ROOT / "thesis" / "Thesis_TDT.md"
OUT_PATH = ROOT / "thesis" / "thesis_final.docx"

# --- Styling Helpers ---

def set_font(run, name="Times New Roman", size=13, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    
    # Ensure Chinese/Vietnamese character rendering support in Word
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rPr.insert(0, rFonts)

def set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6, line_spacing=1.5, first_line_indent=1.0):
    pf = para.paragraph_format
    pf.alignment = align
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)

def set_section_margins(section):
    """TDTU Margins: Top 3.5cm, Bottom 3.0cm, Left 3.5cm, Right 2.0cm"""
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.0)
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)

def configure_page_numbering(section, fmt="decimal", start=None):
    """Manually search and modify or append w:pgNumType elements in CT_SectPr"""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))

def add_page_number(run):
    """Insert Word PAGE field code dynamically"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_formatted_text(para, text, size=13, default_bold=False, default_italic=False):
    """Helper to parse basic inline markdown (**bold** and *italic*)"""
    bold_parts = text.split('**')
    for idx, part in enumerate(bold_parts):
        is_bold = (idx % 2 == 1) or default_bold
        italic_parts = part.split('*')
        for i_idx, sub_part in enumerate(italic_parts):
            is_italic = (i_idx % 2 == 1) or default_italic
            if sub_part:
                run = para.add_run(sub_part)
                set_font(run, size=size, bold=is_bold, italic=is_italic)

def add_table_row(table, values, bold=False, shading=None, center_cols=None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        align = WD_ALIGN_PARAGRAPH.CENTER if (center_cols and i in center_cols) else WD_ALIGN_PARAGRAPH.LEFT
        para.alignment = align
        run = para.add_run(str(val))
        set_font(run, size=11, bold=bold)
        
        # Color Shading
        if shading:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), shading)
            tcPr.append(shd)
    return row

# --- Markdown Block Parser ---

def parse_markdown_block(doc, block):
    block = block.strip()
    if not block:
        return

    # 1. Table Block Detection
    lines = block.split('\n')
    has_table = any(line.strip().startswith('|') for line in lines)
    
    if has_table:
        pre_table_lines = []
        table_lines = []
        for line in lines:
            if line.strip().startswith('|'):
                table_lines.append(line)
            else:
                if not table_lines:  # Only collect lines before the table starts
                    pre_table_lines.append(line)
        
        # Parse lines before the table
        if pre_table_lines:
            pre_text = "\n".join(pre_table_lines).strip()
            if pre_text.startswith("Bảng") or pre_text.startswith("BẢNG"):
                para = doc.add_paragraph()
                run = para.add_run(pre_text)
                set_font(run, size=11, bold=True, italic=True)
                set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=4, first_line_indent=None)
            else:
                parse_markdown_block(doc, pre_text)
                
        # Parse table lines
        if len(table_lines) >= 2:
            headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
            num_cols = len(headers)
            
            t = doc.add_table(rows=0, cols=num_cols)
            t.style = 'Table Grid'
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header Row
            add_table_row(t, headers, bold=True, shading="023047", center_cols=list(range(num_cols)))
            
            # Data Rows
            for row in table_lines[1:]:
                if '---|' in row or '---:|' in row or ':---|' in row:
                    continue
                cells = [c.strip() for c in row.split('|')[1:-1]]
                if len(cells) == num_cols:
                    center_cols = [0]
                    for idx, c in enumerate(cells):
                        if c.replace('%','').replace('$','').replace(',','').replace('.','').replace('-','').strip().isdigit():
                            center_cols.append(idx)
                    add_table_row(t, cells, center_cols=center_cols)
            
            # Add spacing below
            p_space = doc.add_paragraph()
            set_paragraph_format(p_space, space_before=0, space_after=6, first_line_indent=None)
        return

    # 2. Headings
    if block.startswith('# '):
        text = block[2:].strip()
        para = doc.add_paragraph()
        run = para.add_run(text)
        set_font(run, size=14, bold=True)
        if any(keyword in text.upper() for keyword in ["CHƯƠNG", "LỜI CẢM ƠN", "LỜI CAM ĐOAN", "TÓM TẮT", "ABSTRACT", "MỤC LỤC", "DANH MỤC", "TÀI LIỆU THAM KHẢO", "PHỤ LỤC"]):
            set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, space_after=12, first_line_indent=None)
        else:
            set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=12, first_line_indent=None)
        return

    if block.startswith('## '):
        text = block[3:].strip()
        para = doc.add_paragraph()
        run = para.add_run(text)
        set_font(run, size=13, bold=True)
        set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6, first_line_indent=None)
        return

    if block.startswith('### '):
        text = block[4:].strip()
        para = doc.add_paragraph()
        run = para.add_run(text)
        set_font(run, size=13, bold=True, italic=True)
        set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4, first_line_indent=None)
        return

    if block.startswith('#### '):
        text = block[5:].strip()
        para = doc.add_paragraph()
        run = para.add_run(text)
        set_font(run, size=13, italic=True)
        set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=4, first_line_indent=None)
        return

    # 3. Figure Block (Image)
    if block.startswith('!['):
        m = re.match(r'!\[(.*?)\]\((.*?)\)', block)
        if m:
            caption, img_rel_path = m.groups()
            img_abs_path = ROOT / img_rel_path
            
            if not img_abs_path.exists():
                print(f"  [WARN] Image file not found: {img_abs_path}")
                para = doc.add_paragraph()
                run = para.add_run(f"[Hình: {caption} - Đường dẫn: {img_rel_path} không tồn tại]")
                set_font(run, size=11, italic=True)
                set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.CENTER)
                return
            
            # Add image paragraph
            para_img = doc.add_paragraph()
            para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = para_img.add_run()
            run_img.add_picture(str(img_abs_path), width=Cm(14.0))
            set_paragraph_format(para_img, space_before=12, space_after=4, first_line_indent=None)
            
            # Add caption paragraph below
            para_cap = doc.add_paragraph()
            para_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = para_cap.add_run(caption)
            set_font(run_cap, size=11, italic=True)
            set_paragraph_format(para_cap, space_before=0, space_after=12, first_line_indent=None)
        return

    # 4. Bullet List Block
    if block.startswith('* ') or block.startswith('- '):
        lines = block.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('* ') or line.startswith('- '):
                text = line[2:].strip()
                para = doc.add_paragraph(style='List Bullet')
                add_formatted_text(para, text, size=13)
                set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=3, first_line_indent=None)
        return

    # 5. Normal Body Paragraph
    paragraph_text = " ".join(line.strip() for line in block.split('\n') if line.strip())
    para = doc.add_paragraph()
    add_formatted_text(para, paragraph_text, size=13)
    set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6, first_line_indent=1.0)

# --- Document Compiler ---

def compile_document():
    print(f"Reading markdown source from: {MD_PATH}")
    if not MD_PATH.exists():
        print(f"Error: Markdown source file not found at {MD_PATH}")
        sys.exit(1)
        
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        content = f.read()
        
    pages = content.split('<!-- PAGE_BREAK -->')
    print(f"Successfully loaded. Total pages designed: {len(pages)}")
    
    doc = Document()
    
    # ── Section 1: Cover Pages (Pages 1-2) ──
    print("Compiling Section 1 (Covers)...")
    sec1 = doc.sections[0]
    set_section_margins(sec1)
    sec1.header.is_linked_to_previous = False
    sec1.header.paragraphs[0].text = ""
    
    # Parse Cover page (Page 1)
    blocks_p1 = [b.strip() for b in pages[0].split('\n\n') if b.strip()]
    for b in blocks_p1:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        space_before = 0
        if "<br>" in b:
            br_count = b.count("<br>")
            space_before = br_count * 12
            clean_b = b.replace("<br>", "").strip()
        else:
            clean_b = b
            
        add_formatted_text(para, clean_b, size=13 if "TRƯỜNG" not in b and "KHOA" not in b else 14)
        if "#" in clean_b or "**" in clean_b:
            for run in para.runs:
                run.bold = True
                if "#" in clean_b:
                    run.font.size = Pt(16)
        set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=space_before, space_after=6, first_line_indent=None)
    doc.add_page_break()
    
    # Parse Sub-cover (Page 2)
    blocks_p2 = [b.strip() for b in pages[1].split('\n\n') if b.strip()]
    for b in blocks_p2:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        space_before = 0
        if "<br>" in b:
            br_count = b.count("<br>")
            space_before = br_count * 12
            clean_b = b.replace("<br>", "").strip()
        else:
            clean_b = b
            
        add_formatted_text(para, clean_b, size=13 if "TRƯỜNG" not in b and "KHOA" not in b else 14)
        if "#" in clean_b or "**" in clean_b:
            for run in para.runs:
                run.bold = True
                if "#" in clean_b:
                    run.font.size = Pt(16)
        set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=space_before, space_after=6, first_line_indent=None)

    # ── Section 2: Front Matter (Pages 3-10) ──
    print("Compiling Section 2 (Front Matter)...")
    sec2 = doc.add_section()
    set_section_margins(sec2)
    sec2.header.is_linked_to_previous = False
    
    # Configure lower Roman page numbering centered in header
    configure_page_numbering(sec2, fmt="romanLower", start=1) # i
    
    h_para2 = sec2.header.paragraphs[0]
    h_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_para2.text = ""
    run2 = h_para2.add_run()
    set_font(run2, name="Times New Roman", size=13)
    add_page_number(run2)
    
    # Pages 3 to 10 (indices 2 to 9)
    for i in range(2, 10):
        page_text = pages[i]
        blocks = [b.strip() for b in page_text.split('\n\n') if b.strip()]
        for b in blocks:
            parse_markdown_block(doc, b)
        if i < 9:
            doc.add_page_break()

    # ── Section 3: Main Body (Pages 11-42) ──
    print("Compiling Section 3 (Main Body)...")
    sec3 = doc.add_section()
    set_section_margins(sec3)
    sec3.header.is_linked_to_previous = False
    
    # Configure decimal (Arabic) page numbering centered in header restarting at 1
    configure_page_numbering(sec3, fmt="decimal", start=1) # 1
    
    h_para3 = sec3.header.paragraphs[0]
    h_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_para3.text = ""
    run3 = h_para3.add_run()
    set_font(run3, name="Times New Roman", size=13)
    add_page_number(run3)
    
    # Pages 11 to the end (indices 10 to len(pages)-1)
    for i in range(10, len(pages)):
        page_text = pages[i]
        blocks = [b.strip() for b in page_text.split('\n\n') if b.strip()]
        for b in blocks:
            parse_markdown_block(doc, b)
        if i < len(pages) - 1:
            doc.add_page_break()

    # --- Save Document ---
    try:
        doc.save(str(OUT_PATH))
        print(f"\nSuccess! Thesis compiled and saved to: {OUT_PATH}")
        print(f"File size: {OUT_PATH.stat().st_size / 1024:.1f} KB")
    except PermissionError:
        alt_path = ROOT / "thesis" / "thesis_final_full.docx"
        doc.save(str(alt_path))
        print(f"\n[WARN] Permission denied on thesis_final.docx (likely open in Word).")
        print(f"Saved to alternative path: {alt_path}")
        print(f"File size: {alt_path.stat().st_size / 1024:.1f} KB")

if __name__ == "__main__":
    compile_document()
